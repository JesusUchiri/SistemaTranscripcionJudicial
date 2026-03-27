"""
Servicio de exportación de documentos oficiales: DOCX y PDF.
Formato judicial estándar — Poder Judicial del Perú, Distrito de Cusco.

Estructura del documento exportado:
  1. Membrete institucional
  2. Tabla de metadatos del expediente
  3. Cuerpo del acta (generado por IA, editado por digitador)
  4. Bloque de firmas
  5. Footer con número de página
"""
import io
from bs4 import BeautifulSoup, NavigableString, Tag

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from xhtml2pdf import pisa

from app.models.acta import Acta
from app.models.audiencia import Audiencia


# ── Constantes tipográficas ──────────────────────────────────────────
FONT_NAME   = "Times New Roman"
FONT_BODY   = Pt(12)
FONT_SMALL  = Pt(10)
FONT_TITLE  = Pt(13)
LINE_SPACING = Pt(18)  # 1.5 × 12pt


# ── Helpers DOCX ────────────────────────────────────────────────────

def _set_run_font(run, size=None, bold=False, italic=False, underline=False, color=None):
    run.font.name    = FONT_NAME
    run.font.size    = size or FONT_BODY
    run.bold         = bold
    run.italic       = italic
    run.underline    = underline
    if color:
        run.font.color.rgb = color
    # Force font theme compatibility (OOXML)
    r_elem = run._r
    rPr = r_elem.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),     FONT_NAME)
    rFonts.set(qn('w:hAnsi'),     FONT_NAME)
    rFonts.set(qn('w:cs'),        FONT_NAME)
    rFonts.set(qn('w:eastAsia'),  FONT_NAME)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def _set_para_spacing(p, before=0, after=6, line=LINE_SPACING):
    fmt = p.paragraph_format
    fmt.space_before          = Pt(before)
    fmt.space_after           = Pt(after)
    fmt.line_spacing          = line
    fmt.line_spacing_rule     = WD_LINE_SPACING.EXACTLY


def _add_border_para(doc, position='bottom', color='000000', sz='8'):
    """Agrega un párrafo vacío con borde (top o bottom) como separador visual."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    side = OxmlElement(f'w:{position}')
    side.set(qn('w:val'),   'single')
    side.set(qn('w:sz'),    sz)
    side.set(qn('w:space'), '1')
    side.set(qn('w:color'), color)
    pBdr.append(side)
    pPr.append(pBdr)
    return p


def _add_page_number(doc):
    """Footer centrado: 'Página X de Y — Corte Superior de Justicia de Cusco'"""
    section    = doc.sections[0]
    footer     = section.footer
    fp         = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()

    def _add_field(para, field_name):
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.text = field_name
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        r = OxmlElement('w:r')
        r.append(fld_begin)
        r.append(instr)
        r.append(fld_end)
        para._p.append(r)

    for text in ('Corte Superior de Justicia de Cusco   ·   Página ',):
        run = fp.add_run(text)
        _set_run_font(run, size=FONT_SMALL, color=RGBColor(0x66, 0x66, 0x66))

    _add_field(fp, 'PAGE')

    run2 = fp.add_run(' de ')
    _set_run_font(run2, size=FONT_SMALL, color=RGBColor(0x66, 0x66, 0x66))

    _add_field(fp, 'NUMPAGES')


def _add_metadata_table(doc, audiencia: Audiencia, acta: Acta):
    """Tabla de 2 columnas con los metadatos del expediente."""
    expediente = audiencia.expediente or "S/N"
    fecha_str  = _fmt_fecha(audiencia.fecha).upper()
    hora_i     = audiencia.hora_inicio.strftime('%H:%M') if audiencia.hora_inicio else "—"
    hora_f     = audiencia.hora_fin.strftime('%H:%M')    if audiencia.hora_fin    else "—"
    instancia  = (audiencia.instancia or "").upper()
    estado     = (acta.estado or "borrador").upper()

    rows = [
        ("EXPEDIENTE N°",     expediente,                          "ESTADO",    estado),
        ("TIPO DE AUDIENCIA",  (audiencia.tipo_audiencia or "").upper(), "VERSIÓN", f"v{acta.version}"),
        ("INSTANCIA",          instancia,                           "SALA",      audiencia.sala or "—"),
        ("DELITO",             audiencia.delito or "—",             "FECHA",     fecha_str),
        ("IMPUTADO/A",         audiencia.imputado_nombre or "—",    "HORA",      f"{hora_i} – {hora_f}"),
        ("AGRAVIADO/A",        audiencia.agraviado_nombre or "—",   "ESPECIALISTA",
         audiencia.especialista_audiencia or audiencia.especialista_causa or "—"),
    ]

    table = doc.add_table(rows=len(rows), cols=4)
    table.style = 'Table Grid'
    for i, (l1, v1, l2, v2) in enumerate(rows):
        cells = table.rows[i].cells
        for ci, (text, bold) in enumerate([(l1, True), (v1, False), (l2, True), (v2, False)]):
            cell = cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(text)
            run.font.name  = FONT_NAME
            run.font.size  = FONT_SMALL
            run.bold       = bold
            if bold:
                run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(1)
        # Anchos de columna: etiqueta estrecha, valor ancho
        for ci, w in enumerate([Cm(4.5), Cm(6.5), Cm(3.8), Cm(4.0)]):
            cells[ci].width = w

    # Bordes más suaves en la tabla
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'),   'single')
                b.set(qn('w:sz'),    '4')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), 'AAAAAA')
                tcBorders.append(b)
            tcPr.append(tcBorders)

    # Espacio después de la tabla
    p_post = doc.add_paragraph()
    p_post.paragraph_format.space_before = Pt(6)
    p_post.paragraph_format.space_after  = Pt(0)


def _add_signature_block(doc, audiencia: Audiencia):
    """Bloque final de firmas — tres columnas."""
    _add_border_para(doc, position='top', color='888888', sz='4')

    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_note.add_run(
        "Con lo que concluyó la audiencia, siendo las "
        f"{'___' if not audiencia.hora_fin else audiencia.hora_fin.strftime('%H:%M')} horas del mismo día, "
        "suscribiendo los intervinientes en señal de conformidad."
    )
    _set_run_font(r, size=FONT_BODY)
    _set_para_spacing(p_note, before=6, after=18)

    # Tres espacios de firma
    table = doc.add_table(rows=1, cols=3)
    for i, label in enumerate([
        "JUEZ / DIRECTOR DE DEBATES",
        "FISCAL",
        "DEFENSA",
    ]):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()

        # Línea de firma
        p_line = cell.paragraphs[0]
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_line = p_line.add_run("_" * 28)
        _set_run_font(r_line, size=Pt(11))
        _set_para_spacing(p_line, before=30, after=2)

        p_label = cell.add_paragraph()
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_label = p_label.add_run(label)
        _set_run_font(r_label, size=FONT_SMALL, bold=True)
        _set_para_spacing(p_label, before=0, after=0)

        p_cargo = cell.add_paragraph()
        p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cargo = p_cargo.add_run("Firma y Sello")
        _set_run_font(r_cargo, size=Pt(9), color=RGBColor(0x88, 0x88, 0x88))
        _set_para_spacing(p_cargo, before=0, after=0)


# ── Parser HTML → DOCX ──────────────────────────────────────────────

def _inline_runs(p, element):
    """
    Agrega runs al párrafo `p` recorriendo recursivamente los hijos de `element`.
    Preserva negrita, cursiva y subrayado anidados correctamente.
    """
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = p.add_run(text)
                _set_run_font(run)
        elif isinstance(child, Tag):
            tag = child.name.lower() if child.name else ''
            if tag in ('strong', 'b'):
                for gc in child.children:
                    if isinstance(gc, NavigableString):
                        run = p.add_run(str(gc))
                        _set_run_font(run, bold=True)
                    elif isinstance(gc, Tag):
                        run = p.add_run(gc.get_text())
                        _set_run_font(run, bold=True,
                                      italic=(gc.name in ('em', 'i')),
                                      underline=(gc.name == 'u'))
            elif tag in ('em', 'i'):
                for gc in child.children:
                    if isinstance(gc, NavigableString):
                        run = p.add_run(str(gc))
                        _set_run_font(run, italic=True)
                    elif isinstance(gc, Tag):
                        run = p.add_run(gc.get_text())
                        _set_run_font(run, italic=True,
                                      bold=(gc.name in ('strong', 'b')))
            elif tag == 'u':
                run = p.add_run(child.get_text())
                _set_run_font(run, underline=True)
            elif tag == 'br':
                run = p.add_run('\n')
                _set_run_font(run)
            else:
                # Cualquier otro elemento inline — extraer texto plano
                run = p.add_run(child.get_text())
                _set_run_font(run)


def _parse_html_to_docx(doc, html_content: str):
    """
    Convierte el HTML del acta en párrafos Word con tipografía judicial.
    Tags soportados: h1–h3, p, strong/b, em/i, u, br, ul/ol, li, hr.
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    root = soup.body if soup.body else soup

    for element in root.children:
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                _set_run_font(run)
                _set_para_spacing(p)
            continue

        if not isinstance(element, Tag):
            continue

        tag = element.name.lower() if element.name else ''

        # ── Encabezados ─────────────────────────────────────────────
        if tag in ('h1', 'h2', 'h3'):
            level = int(tag[1])
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sizes = {1: FONT_TITLE, 2: Pt(12), 3: Pt(12)}
            p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.line_spacing = Pt(16)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            run = p.add_run(element.get_text(strip=True).upper())
            _set_run_font(run, size=sizes[level], bold=True)

        # ── Párrafo ─────────────────────────────────────────────────
        elif tag == 'p':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _set_para_spacing(p, before=0, after=6)
            _inline_runs(p, element)

        # ── Listas ─────────────────────────────────────────────────
        elif tag in ('ul', 'ol'):
            style = 'List Bullet' if tag == 'ul' else 'List Number'
            for li in element.find_all('li', recursive=False):
                try:
                    p = doc.add_paragraph(style=style)
                except KeyError:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1)
                _set_para_spacing(p, before=0, after=3)
                _inline_runs(p, li)

        # ── Separador horizontal ─────────────────────────────────
        elif tag == 'hr':
            _add_border_para(doc, position='bottom', color='AAAAAA', sz='4')

        # ── Otros elementos con texto ───────────────────────────
        elif element.get_text(strip=True):
            p = doc.add_paragraph()
            _set_para_spacing(p)
            run = p.add_run(element.get_text(strip=True))
            _set_run_font(run)


# ── PDF ─────────────────────────────────────────────────────────────

_MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}

def _fmt_fecha(d) -> str:
    """Formatea fecha como '5 de marzo de 2025' (sin cero inicial, siempre en español)."""
    if not d:
        return "—"
    return f"{d.day} de {_MESES_ES[d.month]} de {d.year}"


def generate_pdf(audiencia: Audiencia, acta: Acta) -> bytes:
    """Genera PDF con formato judicial oficial — A4, Times New Roman, interlineado 1.5."""
    html_content = acta.contenido_editado or acta.contenido_llm or ""
    expediente   = audiencia.expediente or "S/N"
    juzgado      = (audiencia.juzgado or "PODER JUDICIAL DEL PERÚ").upper()
    fecha        = _fmt_fecha(audiencia.fecha).upper()
    tipo         = (audiencia.tipo_audiencia or "AUDIENCIA").upper()
    hora_i       = audiencia.hora_inicio.strftime('%H:%M') if audiencia.hora_inicio else "—"
    hora_f       = audiencia.hora_fin.strftime('%H:%M')    if audiencia.hora_fin    else "—"
    instancia    = (audiencia.instancia or "").upper()
    sala         = audiencia.sala or "—"
    delito       = audiencia.delito or "—"
    imputado     = audiencia.imputado_nombre or "—"
    agraviado    = audiencia.agraviado_nombre or "—"
    especialista = audiencia.especialista_audiencia or audiencia.especialista_causa or "—"
    version      = getattr(acta, 'version', 1)
    estado       = (acta.estado or "borrador").upper()

    # Bloque de firma
    firma_html = """
<table class="firmas-table">
  <tr>
    <td>
      <div class="firma-linea">__________________________</div>
      <div class="firma-cargo">Juez / Director de Debates</div>
      <div class="firma-sub">Firma y Sello</div>
    </td>
    <td>
      <div class="firma-linea">__________________________</div>
      <div class="firma-cargo">Fiscal</div>
      <div class="firma-sub">Firma y Sello</div>
    </td>
    <td>
      <div class="firma-linea">__________________________</div>
      <div class="firma-cargo">Defensa</div>
      <div class="firma-sub">Firma y Sello</div>
    </td>
  </tr>
</table>"""

    watermark = "" if acta.estado == "aprobada" else (
        '<p class="watermark">— DOCUMENTO NO OFICIAL — BORRADOR —</p>'
    )

    template_html = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<title>Acta de Audiencia — Exp. {expediente}</title>
<style>
  @page {{
    size: A4 portrait;
    margin-top: 2.5cm;
    margin-bottom: 3cm;
    margin-left: 3cm;
    margin-right: 2.5cm;
  }}

  /* ── Tipografía base ───────────────────────────────── */
  body {{
    font-family: "Times New Roman", Times, serif;
    font-size: 12pt;
    line-height: 1.5;
    color: #000000;
    text-align: justify;
  }}

  /* ── Membrete ──────────────────────────────────────── */
  .membrete {{
    text-align: center;
    padding-bottom: 10pt;
    border-bottom: 2.5pt solid #000000;
    margin-bottom: 14pt;
  }}
  .membrete .escudo {{
    font-size: 9pt;
    color: #555;
    letter-spacing: 2px;
    text-transform: uppercase;
    margin-bottom: 2pt;
  }}
  .membrete .institucion {{
    font-size: 14pt;
    font-weight: bold;
    text-transform: uppercase;
    letter-spacing: 0.5px;
  }}
  .membrete .juzgado {{
    font-size: 11pt;
    font-weight: bold;
    text-transform: uppercase;
    margin-top: 3pt;
  }}
  .membrete .distrito {{
    font-size: 10pt;
    margin-top: 2pt;
    color: #333;
  }}

  /* ── Tabla de metadatos ─────────────────────────────── */
  .metadatos {{
    width: 100%;
    border-collapse: collapse;
    margin: 10pt 0 14pt 0;
    font-size: 10.5pt;
  }}
  .metadatos td {{
    padding: 3pt 6pt;
    vertical-align: top;
    border: 0.5pt solid #BBBBBB;
  }}
  .metadatos .lbl {{
    font-weight: bold;
    text-transform: uppercase;
    background-color: #F5F5F5;
    width: 22%;
    color: #222;
    white-space: nowrap;
  }}
  .metadatos .val {{
    width: 28%;
  }}

  /* ── Cuerpo del acta ─────────────────────────────────── */
  h3 {{
    font-size: 11.5pt;
    font-weight: bold;
    text-align: center;
    text-transform: uppercase;
    margin: 14pt 0 6pt 0;
    letter-spacing: 0.3px;
  }}
  p {{
    text-align: justify;
    margin: 0 0 6pt 0;
    font-size: 12pt;
    line-height: 1.5;
    text-indent: 0;
  }}
  strong, b {{
    font-weight: bold;
  }}
  em, i {{
    font-style: italic;
  }}
  ul, ol {{
    margin: 4pt 0 6pt 18pt;
    padding: 0;
  }}
  li {{
    margin: 2pt 0;
    font-size: 12pt;
    line-height: 1.5;
  }}

  /* ── Watermark borrador ──────────────────────────────── */
  .watermark {{
    font-size: 48pt;
    color: #DDDDDD;
    font-weight: bold;
    text-transform: uppercase;
    text-align: center;
    letter-spacing: 8px;
    margin: 8pt 0;
  }}

  /* ── Bloque de firmas (tabla HTML) ───────────────────── */
  .firmas-table {{
    width: 100%;
    margin-top: 36pt;
    border-top: 1pt solid #888888;
    padding-top: 4pt;
    border-collapse: collapse;
  }}
  .firmas-table td {{
    width: 33%;
    text-align: center;
    padding: 24pt 8pt 4pt 8pt;
    vertical-align: top;
    border: none;
  }}
  .firma-linea {{
    font-size: 11pt;
    margin-bottom: 4pt;
  }}
  .firma-cargo {{
    font-size: 9.5pt;
    font-weight: bold;
    text-transform: uppercase;
  }}
  .firma-sub {{
    font-size: 9pt;
    color: #777777;
    margin-top: 2pt;
  }}

  /* ── Footer ──────────────────────────────────────────── */
  .pie-pagina {{
    position: fixed;
    bottom: 0;
    left: 0;
    right: 0;
    text-align: center;
    font-size: 8.5pt;
    color: #777777;
    border-top: 0.5pt solid #CCCCCC;
    padding-top: 4pt;
  }}
</style>
</head>
<body>

<!-- ══ MEMBRETE INSTITUCIONAL ══════════════════════════════════ -->
<div class="membrete">
  <div class="escudo">⚖ Poder Judicial del Perú</div>
  <div class="institucion">Corte Superior de Justicia de Cusco</div>
  <div class="juzgado">{juzgado}</div>
  <div class="distrito">Distrito Judicial de Cusco</div>
</div>

<!-- ══ TABLA DE METADATOS ═══════════════════════════════════════ -->
<table class="metadatos">
  <tr>
    <td class="lbl">Expediente N°</td>
    <td class="val">{expediente}</td>
    <td class="lbl">Estado</td>
    <td class="val">{estado}</td>
  </tr>
  <tr>
    <td class="lbl">Tipo de Audiencia</td>
    <td class="val">{tipo}</td>
    <td class="lbl">Versión</td>
    <td class="val">v{version}</td>
  </tr>
  <tr>
    <td class="lbl">Instancia</td>
    <td class="val">{instancia}</td>
    <td class="lbl">Sala</td>
    <td class="val">{sala}</td>
  </tr>
  <tr>
    <td class="lbl">Delito</td>
    <td class="val">{delito}</td>
    <td class="lbl">Fecha</td>
    <td class="val">{fecha}</td>
  </tr>
  <tr>
    <td class="lbl">Imputado/a</td>
    <td class="val">{imputado}</td>
    <td class="lbl">Hora</td>
    <td class="val">{hora_i} – {hora_f}</td>
  </tr>
  <tr>
    <td class="lbl">Agraviado/a</td>
    <td class="val">{agraviado}</td>
    <td class="lbl">Especialista</td>
    <td class="val">{especialista}</td>
  </tr>
</table>

<!-- ══ AVISO BORRADOR (si aplica) ══════════════════════════════ -->
{watermark}

<!-- ══ CUERPO DEL ACTA (generado por IA / editado) ════════════ -->
{html_content}

<!-- ══ BLOQUE DE FIRMAS ════════════════════════════════════════ -->
{firma_html}

<!-- ══ PIE DE PÁGINA ══════════════════════════════════════════ -->
<div class="pie-pagina">
  Corte Superior de Justicia de Cusco &nbsp;·&nbsp; Exp. {expediente} &nbsp;·&nbsp;
  Página <pdf:pagenumber> de <pdf:pagecount>
</div>

</body>
</html>"""

    buf = io.BytesIO()
    status = pisa.CreatePDF(io.StringIO(template_html), dest=buf)
    if status.err:
        raise RuntimeError(f"xhtml2pdf error al generar PDF (código {status.err})")
    return buf.getvalue()


# ── DOCX ─────────────────────────────────────────────────────────────

def generate_docx(audiencia: Audiencia, acta: Acta) -> bytes:
    """Genera DOCX con formato judicial oficial — A4, Times New Roman 12pt, interlineado 1.5."""
    doc = DocxDocument()

    # ── Página A4 con márgenes judiciales ──────────────────────────
    section               = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(3.0)

    # ── Estilos base ────────────────────────────────────────────────
    normal = doc.styles['Normal']
    normal.font.name = FONT_NAME
    normal.font.size = FONT_BODY
    normal.paragraph_format.line_spacing      = LINE_SPACING
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.space_after       = Pt(6)

    # ── Encabezado institucional ─────────────────────────────────
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after  = Pt(0)

    r1 = hp.add_run("CORTE SUPERIOR DE JUSTICIA DE CUSCO")
    _set_run_font(r1, size=Pt(11), bold=True)
    hp.add_run("\n")
    r2 = hp.add_run((audiencia.juzgado or "PODER JUDICIAL DEL PERÚ").upper())
    _set_run_font(r2, size=Pt(10))
    hp.add_run("\n")
    r3 = hp.add_run("DISTRITO JUDICIAL DE CUSCO")
    _set_run_font(r3, size=Pt(9), color=RGBColor(0x44, 0x44, 0x44))

    # Línea bajo el encabezado
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Tabla de metadatos ──────────────────────────────────────
    _add_metadata_table(doc, audiencia, acta)

    # ── Separador superior del cuerpo ───────────────────────────
    _add_border_para(doc, position='bottom', color='444444', sz='6')

    # ── Contenido del acta (HTML → Word) ────────────────────────
    html_content = acta.contenido_editado or acta.contenido_llm or ""
    _parse_html_to_docx(doc, html_content)

    # ── Bloque de firmas ─────────────────────────────────────────
    _add_signature_block(doc, audiencia)

    # ── Footer con número de página ──────────────────────────────
    _add_page_number(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
