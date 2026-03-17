import pytest
from app.services.export_service import generate_pdf, generate_docx
from app.models.acta import Acta
from app.models.audiencia import Audiencia
import io

def test_generate_pdf():
    """Prueba la generación de PDF desde HTML de Acta (Sprint 10)."""
    aud = Audiencia(expediente="MOCK-001", juzgado="Juzgado Penal")
    acta = Acta(contenido_editado="<p>Texto <b>Oficial</b> del <i>Acta</i></p>")
    
    pdf_bytes = generate_pdf(aud, acta)
    assert pdf_bytes is not None
    assert isinstance(pdf_bytes, bytes)
    assert pdf_bytes.startswith(b"%PDF-") # Magic header for pdffiles

def test_generate_docx():
    """Prueba la compilación del DOCX nativo usando beautifulsoup/python-docx (Sprint 10)."""
    aud = Audiencia(expediente="MOCK-002", juzgado="Juzgado Civil")
    acta = Acta(contenido_editado="<h2>TÍTULO PRINCIPAL</h2><p>Texto <b>Oficial</b> del Acta Civil</p>")
    
    docx_bytes = generate_docx(aud, acta)
    
    assert docx_bytes is not None
    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 1000 # OLE zip file bytes size
    
    from docx import Document
    
    doc = Document(io.BytesIO(docx_bytes))
    
    # Check elements parsing inside document object
    assert len(doc.paragraphs) > 0
    header_found = False
    for section in doc.sections:
        for para in section.header.paragraphs:
            if "Juzgado Civil" in para.text:
                header_found = True
    assert header_found, "Header from audiencia properties not injected properly"
